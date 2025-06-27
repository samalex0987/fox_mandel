from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

app = FastAPI()

# Assuming you have a `processing_status` dictionary somewhere in your application
processing_status = {}

@app.get("/download/{session_id}/{file_type}")
async def download_file(session_id: str, file_type: str):
    """Download generated report file (Markdown or DOCX)"""

    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    status_data = processing_status[session_id]
    final_output = status_data.get("final_output")

    if not final_output:
        raise HTTPException(status_code=404, detail="No content available for export")

    # Create output folder
    output_dir = os.path.join("outputs", session_id)
    os.makedirs(output_dir, exist_ok=True)

    # Common logo path
    logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/foxmandal-logo.png"))

    if file_type == "markdown":
        markdown_path = os.path.join(output_dir, "report.md")

        # Compose markdown content with logo
        if os.path.exists(logo_path):
            logo_markdown = f"""# FOXMandal

![FOXMandal Logo]({logo_path})

---

"""
        else:
            logo_markdown = "# FOXMandal\n\n---\n\n"

        content_with_logo = logo_markdown + final_output

        # Save to file
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(content_with_logo)

        status_data["markdown_path"] = markdown_path

        return FileResponse(
            path=markdown_path,
            media_type="text/markdown",
            filename="report.md",
            content_disposition_type="attachment"
        )

    elif file_type == "docx":
        docx_path = os.path.join(output_dir, "report.docx")

        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # Add logo (resizable)
            if os.path.exists(logo_path):
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(2))  # Resize image here

                    # Add blue text below the image
                    paragraph = doc.add_paragraph()
                    blue_run = paragraph.add_run("FOXMandal")  # The text you want to be blue
                    blue_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
                except Exception as img_error:
                    print(f"Logo image insertion failed: {img_error}")

            # Title
            doc.add_heading('FOXMandal', 0)
            doc.add_paragraph("---")

            # Process markdown-style content
            lines = final_output.split('\n')
            current_table = None

            for line in lines:
                line_stripped = line.strip()

                if not line_stripped:
                    doc.add_paragraph()
                    continue

                if line_stripped.startswith('### '):
                    doc.add_heading(line_stripped[4:], level=3)
                elif line_stripped.startswith('## '):
                    doc.add_heading(line_stripped[3:], level=2)
                elif line_stripped.startswith('# '):
                    doc.add_heading(line_stripped[2:], level=1)

                elif line_stripped.startswith('| ') and line_stripped.endswith(' |'):
                    cells = [cell.strip() for cell in line_stripped.strip('|').split('|')]
                    if current_table is None:
                        current_table = doc.add_table(rows=1, cols=len(cells))
                        current_table.style = 'Table Grid'
                        for i, cell_text in enumerate(cells):
                            cell = current_table.rows[0].cells[i]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                    else:
                        row = current_table.add_row()
                        for i, cell_text in enumerate(cells):
                            if i < len(row.cells):
                                row.cells[i].text = cell_text

                elif re.match(r'^\|[\s\-\|]+\|$', line_stripped):
                    continue  # Skip table separator lines

                elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                    current_table = None
                    doc.add_paragraph(line_stripped[2:], style='List Bullet')

                elif re.match(r'^\d+\.\s', line_stripped):
                    current_table = None
                    doc.add_paragraph(line_stripped[line_stripped.find('.') + 1:].strip(), style='List Number')

                elif '**' in line_stripped:
                    current_table = None
                    paragraph = doc.add_paragraph()
                    parts = line_stripped.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            paragraph.add_run(part)
                        else:
                            paragraph.add_run(part).bold = True

                else:
                    current_table = None
                    doc.add_paragraph(line_stripped)

            # Save DOCX
            doc.save(docx_path)
            status_data["docx_path"] = docx_path

            return FileResponse(
                path=docx_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="report.docx",
                content_disposition_type="attachment"
            )

        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to generate DOCX: {str(e)}"
            )

    else:
        raise HTTPException(status_code=400, detail="Invalid file type requested")
