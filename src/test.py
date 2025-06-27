@app.get("/download/{session_id}/{file_type}")
async def download_file(session_id: str, file_type: str):
    """Download generated report file"""
    if session_id not in processing_status:
        raise HTTPException(status_code=404, detail="Processing session not found")
    
    status_data = processing_status[session_id]
    
    if file_type == "markdown":
        # Get the final output directly from status data
        final_output = status_data.get("final_output")
        if not final_output:
            raise HTTPException(status_code=404, detail="No markdown content available")
        
        # Create output directory
        output_dir = os.path.join("outputs", session_id)
        os.makedirs(output_dir, exist_ok=True)
        markdown_path = os.path.join(output_dir, "report.md")
        
        # Add logo at the beginning of the document with proper path
        logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/foxmandal-logo.png"))
        
        # Check if logo file exists
        if os.path.exists(logo_path):
            logo_markdown = f"""# FOXMandal

![FOXMandal Logo]({logo_path})

---

"""
        else:
            # Fallback if logo doesn't exist
            logo_markdown = """# FOXMandal

---

"""
        
        # Combine logo with the final output
        content_with_logo = logo_markdown + final_output
        
        # Write the content to the file
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(content_with_logo)
        
        # Update the status with the markdown path
        status_data["markdown_path"] = markdown_path
        
        # Return the file
        return FileResponse(
            path=markdown_path,
            media_type="text/markdown",
            filename="report.md",
            content_disposition_type="attachment"
        )
        
    elif file_type == "docx":
        # First ensure we have the markdown content
        final_output = status_data.get("final_output")
        if not final_output:
            raise HTTPException(status_code=404, detail="No content available for conversion")
        
        # Create output directory
        output_dir = os.path.join("outputs", session_id)
        os.makedirs(output_dir, exist_ok=True)
        docx_path = os.path.join(output_dir, "report.docx")
        
        try:
            # Method 1: Try pypandoc first
            try:
                # Ensure pandoc is available
                try:
                    pypandoc.download_pandoc()
                except:
                    pass
                
                # Create a temporary markdown file with logo
                temp_dir = os.path.join("temp", session_id)
                os.makedirs(temp_dir, exist_ok=True)
                temp_md_path = os.path.join(temp_dir, "temp_report.md")
                
                # Add logo at the beginning of the document
                logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/foxmandal-logo.png"))
                
                if os.path.exists(logo_path):
                    logo_markdown = f"""# FOXMandal

![FOXMandal Logo]({logo_path})

---

"""
                else:
                    logo_markdown = """# FOXMandal

---

"""
                
                # Combine logo with the final output
                content_with_logo = logo_markdown + final_output
                
                # Write the content to temporary markdown file
                with open(temp_md_path, 'w', encoding='utf-8') as f:
                    f.write(content_with_logo)
                
                # Convert markdown to DOCX with improved formatting
                extra_args = [
                    '--standalone',
                    '--toc',
                    '--toc-depth=3',
                    '--highlight-style=tango',
                    '--wrap=none',
                    '--dpi=300'
                ]
                
                # Add reference doc if template exists
                template_path = os.path.join(os.path.dirname(__file__), "template.docx")
                if os.path.exists(template_path):
                    extra_args.append(f'--reference-doc={template_path}')
                
                pypandoc.convert_file(
                    temp_md_path,
                    'docx',
                    outputfile=docx_path,
                    extra_args=extra_args
                )
                
                # Clean up temporary file
                os.remove(temp_md_path)
                
                # Update the status
                status_data["docx_path"] = docx_path
                
                # Return the file
                return FileResponse(
                    path=docx_path,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename="report.docx",
                    content_disposition_type="attachment"
                )
                
            except Exception as pandoc_error:
                print(f"Pandoc conversion failed: {pandoc_error}")
                raise pandoc_error
                
        except Exception as e:
            # Method 2: If pandoc conversion fails, try python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import re
                
                # Create a new Document
                doc = Document()
                
                # Set default font and size
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(11)
                
                # Add logo at the top if it exists
                logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/foxmandal-logo.png"))
                if os.path.exists(logo_path):
                    try:
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.add_picture(logo_path, width=Inches(2))
                    except Exception as img_error:
                        print(f"Failed to add logo image: {img_error}")
                
                # Add title
                doc.add_heading('FOXMandal', 0)
                doc.add_paragraph("---")
                
                # Process the final output content
                lines = final_output.split('\n')
                current_table = None
                table_headers = []
                
                for line in lines:
                    line_stripped = line.strip()
                    
                    if not line_stripped:
                        # Add empty paragraph for spacing
                        doc.add_paragraph()
                        continue
                    
                    # Handle headers
                    if line_stripped.startswith('### '):
                        doc.add_heading(line_stripped[4:], level=3)
                    elif line_stripped.startswith('## '):
                        doc.add_heading(line_stripped[3:], level=2)
                    elif line_stripped.startswith('# '):
                        doc.add_heading(line_stripped[2:], level=1)
                    
                    # Handle tables
                    elif line_stripped.startswith('| ') and line_stripped.endswith(' |'):
                        cells = [cell.strip() for cell in line_stripped.strip('|').split('|')]
                        
                        if current_table is None:
                            # Create new table
                            current_table = doc.add_table(rows=1, cols=len(cells))
                            current_table.style = 'Table Grid'
                            table_headers = cells
                            
                            # Add header row
                            header_row = current_table.rows[0]
                            for i, cell_text in enumerate(cells):
                                header_row.cells[i].text = cell_text
                                # Make header bold
                                for paragraph in header_row.cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                        else:
                            # Add data row
                            row = current_table.add_row()
                            for i, cell_text in enumerate(cells):
                                if i < len(row.cells):
                                    row.cells[i].text = cell_text
                    
                    # Handle table separator (---|---|---)
                    elif re.match(r'^\|[\s\-\|]+\|$', line_stripped):
                        continue  # Skip table separator lines
                    
                    # Handle bullet points
                    elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                        current_table = None  # End any current table
                        doc.add_paragraph(line_stripped[2:], style='List Bullet')
                    
                    # Handle numbered lists
                    elif re.match(r'^\d+\.\s', line_stripped):
                        current_table = None  # End any current table
                        doc.add_paragraph(line_stripped[line_stripped.find('.') + 1:].strip(), style='List Number')
                    
                    # Handle bold text **text**
                    elif '**' in line_stripped:
                        current_table = None  # End any current table
                        paragraph = doc.add_paragraph()
                        parts = line_stripped.split('**')
                        for i, part in enumerate(parts):
                            if i % 2 == 0:
                                paragraph.add_run(part)
                            else:
                                paragraph.add_run(part).bold = True
                    
                    # Regular paragraph
                    else:
                        current_table = None  # End any current table
                        doc.add_paragraph(line_stripped)
                
                # Save the document
                doc.save(docx_path)
                
                # Update the status
                status_data["docx_path"] = docx_path
                
                # Return the file
                return FileResponse(
                    path=docx_path,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename="report.docx",
                    content_disposition_type="attachment"
                )
                
            except Exception as docx_error:
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to generate DOCX with pandoc: {str(e)}. Alternative method also failed: {str(docx_error)}"
                )
    else:
        raise HTTPException(status_code=400, detail="Invalid file type requested")
